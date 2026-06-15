"""
Extract the content of documents the Host uploads to a meeting, to insert into
the minutes for the sub-agents to read.

Supports:
- PDF  (.pdf)          -> pypdf
- Word (.docx)         -> python-docx
- Excel (.xlsx)        -> openpyxl
- Text (.txt/.csv/.md/.json/code...) -> read directly
- Images (.png/.jpg/...)  -> NO text extraction; the agent uses the image-reading
                          tool (Read) on the path to "view" it.
"""

import os

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif"}

TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".log", ".srt", ".vtt",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".htm", ".css", ".xml",
    ".yaml", ".yml", ".ini", ".cfg", ".toml", ".sql", ".sh", ".bat", ".ps1",
}

# Truncate to avoid bloating the context. The full version stays on disk -> the agent can
# use the file-reading tool to open it in full if needed.
MAX_CHARS = 20000


def is_image(filename: str) -> bool:
    return os.path.splitext(filename)[1].lower() in IMAGE_EXTS


def _truncate(text: str, path: str) -> str:
    text = (text or "").strip()
    if len(text) <= MAX_CHARS:
        return text
    return (
        text[:MAX_CHARS]
        + f"\n\n…[TRUNCATED — long document. The full version is at: {path} "
        "— use the file-reading tool (Read) to open this path and see it in full if needed.]"
    )


def extract_text(path: str) -> str:
    """Return the document's text content (truncated if too long)."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            return _truncate(_pdf(path), path)
        if ext == ".docx":
            return _truncate(_docx(path), path)
        if ext == ".xlsx":
            return _truncate(_xlsx(path), path)
        if ext in TEXT_EXTS or ext == "":
            return _truncate(_plain(path), path)
    except Exception as e:
        return (
            f"[Could not read the file content ({type(e).__name__}: {e}). "
            f"The agent can try the file-reading tool at: {path}]"
        )
    # unknown format: let the agent open it with a tool if it wants
    return (
        f"[Format '{ext}' is not auto-extracted. "
        f"The agent should use the file-reading tool at {path} if it needs the content.]"
    )


def _plain(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def _pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages, 1):
        t = (page.extract_text() or "").strip()
        if t:
            pages.append(f"--- Page {i} ---\n{t}")
    out = "\n\n".join(pages).strip()
    return out or "[Could not extract text from the PDF — it may be a scanned/image PDF. The agent can read the PDF file directly with a tool if needed.]"


def _docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for tb in doc.tables:                       # also collect table content
        for row in tb.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip() or "[Empty Word file.]"


def _xlsx(path: str) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    out = []
    try:
        for ws in wb.worksheets:
            out.append(f"=== Sheet: {ws.title} ===")
            n = 0
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(c.strip() for c in cells):
                    out.append(" | ".join(cells))
                    n += 1
                if n >= 200:                    # cap at 200 non-empty rows per sheet
                    out.append("…[more — capped at 200 rows for this sheet]")
                    break
    finally:
        wb.close()
    return "\n".join(out).strip() or "[Empty Excel file.]"
