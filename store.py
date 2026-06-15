"""
Storage:
- PROJECT: each project = 1 JSON file in projects/. Each project is its own
  WORKSPACE: its own meetings + its own memory + its own documents.
- MEETING: each meeting = 1 JSON file in meetings/ (with a project_id field).
- MEMORY: per project -> memory/<project_id>/<key>.md (+ _team.md).
Plus exporting the minutes to a Word file (.docx).
"""

import os
import re
import json
import uuid
import shutil
import datetime

from docx import Document

import roster

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MEETINGS_DIR = os.path.join(BASE_DIR, "meetings")
UPLOADS_DIR = os.path.join(MEETINGS_DIR, "uploads")
PROJECTS_DIR = os.path.join(BASE_DIR, "projects")
MEMORY_DIR = os.path.join(BASE_DIR, "memory")
os.makedirs(MEETINGS_DIR, exist_ok=True)

TEAM_KEY = "_team"                  # special key for the team's SHARED memory
DEFAULT_PROJECT_ID = "prj_default"  # default project id (holds data migrated from older versions)


def _path(mid: str) -> str:
    return os.path.join(MEETINGS_DIR, f"{mid}.json")


def new_id() -> str:
    return "mtg_" + uuid.uuid4().hex[:12]


def now_iso() -> str:
    return datetime.datetime.now().isoformat(timespec="seconds")


def save(data: dict) -> None:
    data["updated_at"] = now_iso()
    with open(_path(data["id"]), "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load(mid: str):
    p = _path(mid)
    if not os.path.isfile(p):
        return None
    with open(p, encoding="utf-8") as f:
        return json.load(f)


def list_all(project_id: str = None) -> list:
    """List of meetings. If project_id is given -> only meetings in that project."""
    items = []
    for fn in os.listdir(MEETINGS_DIR):
        if not fn.endswith(".json"):
            continue
        try:
            with open(os.path.join(MEETINGS_DIR, fn), encoding="utf-8") as f:
                d = json.load(f)
        except Exception:
            continue
        if project_id is not None and d.get("project_id") != project_id:
            continue
        items.append({
            "id": d.get("id"),
            "title": d.get("title") or "(untitled)",
            "project_id": d.get("project_id"),
            "created_at": d.get("created_at"),
            "updated_at": d.get("updated_at"),
            "agent_keys": d.get("agent_keys", []),
            "count": len([e for e in d.get("log", []) if e.get("speaker")]),
        })
    items.sort(key=lambda x: x.get("updated_at") or "", reverse=True)
    return items


def delete(mid: str) -> None:
    for ext in (".json", ".docx"):
        p = os.path.join(MEETINGS_DIR, f"{mid}{ext}")
        if os.path.isfile(p):
            os.remove(p)
    folder = os.path.join(UPLOADS_DIR, mid)            # also delete any uploaded documents
    if os.path.isdir(folder):
        shutil.rmtree(folder, ignore_errors=True)


# ---------- PROJECTS ----------

def _project_path(pid: str) -> str:
    return os.path.join(PROJECTS_DIR, f"{pid}.json")


def new_project_id() -> str:
    return "prj_" + uuid.uuid4().hex[:12]


def save_project(d: dict) -> None:
    os.makedirs(PROJECTS_DIR, exist_ok=True)
    d["updated_at"] = now_iso()
    with open(_project_path(d["id"]), "w", encoding="utf-8") as f:
        json.dump(d, f, ensure_ascii=False, indent=2)


def load_project(pid: str):
    p = _project_path(pid)
    if not pid or not os.path.isfile(p):
        return None
    with open(p, encoding="utf-8") as f:
        return json.load(f)


def list_projects() -> list:
    os.makedirs(PROJECTS_DIR, exist_ok=True)
    items = []
    for fn in os.listdir(PROJECTS_DIR):
        if not fn.endswith(".json"):
            continue
        try:
            with open(os.path.join(PROJECTS_DIR, fn), encoding="utf-8") as f:
                d = json.load(f)
        except Exception:
            continue
        items.append({
            "id": d.get("id"),
            "name": d.get("name") or "(unnamed)",
            "created_at": d.get("created_at"),
            "updated_at": d.get("updated_at"),
            "count": len(list_all(d.get("id"))),
        })
    items.sort(key=lambda x: x.get("updated_at") or "", reverse=True)
    return items


def create_project(name: str) -> dict:
    d = {"id": new_project_id(), "name": (name or "New project").strip(), "created_at": now_iso()}
    save_project(d)
    return d


def delete_project(pid: str) -> None:
    """Delete the project + all of its meetings + all of the project's memory."""
    for m in list_all(pid):
        delete(m["id"])
    mdir = os.path.join(MEMORY_DIR, pid)
    if os.path.isdir(mdir):
        shutil.rmtree(mdir, ignore_errors=True)
    p = _project_path(pid)
    if os.path.isfile(p):
        os.remove(p)


def ensure_default() -> str:
    """Make sure there is always at least 1 project; return the id of a valid project."""
    projs = list_projects()
    if projs:
        return projs[0]["id"]
    return create_project("General project")["id"]


# ---------- documents the Host uploads to a meeting ----------

def _safe_name(name: str) -> str:
    """Sanitize a file name (guard against path traversal and odd characters)."""
    name = os.path.basename(name or "file")
    name = re.sub(r'[^\w.\-() ]+', "_", name, flags=re.UNICODE).strip()
    return name or "file"


def uploads_dir(mid: str) -> str:
    d = os.path.join(UPLOADS_DIR, mid)
    os.makedirs(d, exist_ok=True)
    return d


def save_upload(mid: str, filename: str, raw: bytes):
    """Save an uploaded file to meetings/uploads/<mid>/. Returns (path, safe_name)."""
    safe = _safe_name(filename)
    folder = uploads_dir(mid)
    path = os.path.join(folder, safe)
    base, ext = os.path.splitext(safe)
    i = 1
    while os.path.exists(path):                         # avoid overwriting on name clash
        safe = f"{base}({i}){ext}"
        path = os.path.join(folder, safe)
        i += 1
    with open(path, "wb") as f:
        f.write(raw)
    return path, safe


def upload_path(mid: str, name: str):
    """Path of an uploaded file (sanitized name) — for serving downloads/display."""
    return os.path.join(uploads_dir(mid), _safe_name(name))


# ---------- memory (per PROJECT: each agent's own + the team's shared) ----------

def _mem_owner(key: str) -> str:
    return "Whole team" if key == TEAM_KEY else roster.name_of(key)


def memory_dir(project_id: str) -> str:
    d = os.path.join(MEMORY_DIR, project_id or DEFAULT_PROJECT_ID)
    os.makedirs(d, exist_ok=True)
    return d


def memory_path(project_id: str, key: str) -> str:
    return os.path.join(memory_dir(project_id), f"{_safe_name(key)}.md")


def read_memory(project_id: str, key: str) -> str:
    p = memory_path(project_id, key)
    if os.path.isfile(p):
        with open(p, encoding="utf-8") as f:
            return f.read().strip()
    return ""


def append_memory(project_id: str, key: str, fact: str) -> bool:
    """Append one thing to remember. Skip if empty/duplicate. True if written."""
    fact = " ".join((fact or "").split()).strip()
    if not fact:
        return False
    existing = read_memory(project_id, key)
    if fact.lower() in existing.lower():        # avoid remembering an exact duplicate
        return False
    date = datetime.date.today().isoformat()
    p = memory_path(project_id, key)
    with open(p, "a", encoding="utf-8") as f:
        if not existing:
            f.write(f"# Memory of {_mem_owner(key)}\n\n")
        f.write(f"- ({date}) {fact}\n")
    return True


def clear_memory(project_id: str, key: str) -> None:
    p = memory_path(project_id, key)
    if os.path.isfile(p):
        os.remove(p)


def all_memories(project_id: str) -> dict:
    return {k: read_memory(project_id, k) for k in roster.keys()}


def memory_facts(project_id: str, key: str) -> list:
    """List of 'fact' lines (starting with '- ') in the memory."""
    return [ln for ln in read_memory(project_id, key).splitlines() if ln.strip().startswith("- ")]


def memory_fact_count(project_id: str, key: str) -> int:
    return len(memory_facts(project_id, key))


def memory_for_prompt(project_id: str, key: str, max_chars: int) -> str:
    """The 'facts' to load into the prompt (memory lines ONLY). Trim to the MOST RECENT if over the limit."""
    facts = memory_facts(project_id, key)
    if not facts:
        return ""
    kept, total, truncated = [], 0, False
    for ln in reversed(facts):              # prefer the newest facts (at the end)
        if kept and total + len(ln) + 1 > max_chars:
            truncated = True
            break
        kept.append(ln)
        total += len(ln) + 1
    kept.reverse()
    prefix = "(showing only the most recent memories)\n" if truncated else ""
    return prefix + "\n".join(kept)


def write_memory(project_id: str, key: str, facts: list) -> int:
    """OVERWRITE the memory with a cleaned list of facts (used after CONSOLIDATING). Returns the count."""
    facts = [" ".join(f.split()).strip() for f in facts if f and f.strip()]
    p = memory_path(project_id, key)
    if not facts:
        if os.path.isfile(p):
            os.remove(p)
        return 0
    with open(p, "w", encoding="utf-8") as f:
        f.write(f"# Memory of {_mem_owner(key)}\n\n")
        for fa in facts:
            fa = fa.lstrip("-• ").strip()
            f.write(f"- {fa}\n")
    return len(facts)


def export_docx(d: dict) -> str:
    """Build the .docx minutes, return the file path."""
    doc = Document()
    doc.add_heading(d.get("title") or "Meeting minutes", level=0)

    p = doc.add_paragraph()
    p.add_run("Created at: ").bold = True
    p.add_run(str(d.get("created_at", "")))

    parts = [roster.name_of(k) for k in d.get("agent_keys", [])]
    p2 = doc.add_paragraph()
    p2.add_run("Participants: ").bold = True
    p2.add_run(", ".join(parts))

    doc.add_paragraph("")
    doc.add_heading("Content", level=1)

    for e in d.get("log", []):
        speaker = e.get("speaker", "")
        text = e.get("text", "")
        para = doc.add_paragraph()
        r = para.add_run(f"{speaker}: ")
        r.bold = True
        para.add_run(text)

    out = os.path.join(MEETINGS_DIR, f"{d['id']}.docx")
    doc.save(out)
    return out


# ---------- MIGRATION of data from older versions (runs once, idempotent) ----------

def ensure_seed() -> None:
    """First upgrade to 'Projects': create the default project and gather all old data
    (meetings without a project_id + memory in root-level memory/*.md) into it. Loses nothing."""
    os.makedirs(PROJECTS_DIR, exist_ok=True)
    if list_projects():
        return  # already has a project -> already migrated

    # 1) Create the default project with a fixed id
    save_project({"id": DEFAULT_PROJECT_ID, "name": "General project", "created_at": now_iso()})

    # 2) Migrate old memory: memory/*.md (root level) -> memory/<DEFAULT>/
    if os.path.isdir(MEMORY_DIR):
        dest = os.path.join(MEMORY_DIR, DEFAULT_PROJECT_ID)
        os.makedirs(dest, exist_ok=True)
        for fn in os.listdir(MEMORY_DIR):
            src = os.path.join(MEMORY_DIR, fn)
            if os.path.isfile(src) and fn.endswith(".md"):
                target = os.path.join(dest, fn)
                if not os.path.exists(target):
                    shutil.move(src, target)

    # 3) Assign every old meeting (without a project_id) to the default project
    for fn in os.listdir(MEETINGS_DIR):
        if not fn.endswith(".json"):
            continue
        fp = os.path.join(MEETINGS_DIR, fn)
        try:
            with open(fp, encoding="utf-8") as f:
                d = json.load(f)
        except Exception:
            continue
        if not d.get("project_id"):
            d["project_id"] = DEFAULT_PROJECT_ID
            with open(fp, "w", encoding="utf-8") as f:
                json.dump(d, f, ensure_ascii=False, indent=2)


ensure_seed()
